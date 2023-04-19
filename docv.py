
import subprocess
import os
import sys
import subprocess
# 检查并安装所需库 python docv.py --package
def package_with_nuitka(script_name):
    try:
        output = subprocess.check_output(["nuitka", "--version"])
    except FileNotFoundError:
        print("Nuitka not found. Please install Nuitka and try again.")
        sys.exit(1)

    command = [
        sys.executable,
        "-m",
        "nuitka",
        "--onefile",
        "--windows-disable-console",
        "--enable-plugin=tk-inter",
        script_name,
    ]

    subprocess.run(command)

if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--package":
        script_name = os.path.basename(sys.argv[0])
        package_with_nuitka(script_name)
    else:
        # Your main program code goes here
        print("Running main program...")
def install_package(package):
   subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
   import docx
except ImportError:
   print("正在安装python-docx库...")
   install_package("python-docx")
   import docx
try:
   import tkinter as tk
   from tkinter import filedialog
   from tkinter import messagebox
except ImportError:
   print("正在安装Tkinter库...")
   install_package("tk")
   import tkinter as tk
   from tkinter import filedialog
   from tkinter import messagebox
# 复制格式的函数和应用模板格式的函数（与之前的示例相同）
def copy_formatting(src_paragraph, dest_paragraph):
   dest_paragraph.style = src_paragraph.style
   for run, dest_run in zip(src_paragraph.runs, dest_paragraph.runs):
       dest_run.font.name = run.font.name
       dest_run.font.size = run.font.size
       dest_run.bold = run.bold
       dest_run.italic = run.italic
       dest_run.underline = run.underline
       dest_run.font.color.rgb = run.font.color.rgb
def apply_template_formatting(template_doc, target_doc):
   template_styles = {}
   for para in template_doc.paragraphs:
       style_name = para.style.name
       if style_name not in template_styles:
           template_styles[style_name] = para
   for para in target_doc.paragraphs:
       style_name = para.style.name
       if style_name in template_styles:
           copy_formatting(template_styles[style_name], para)
# GUI相关代码
def browse_template_file():
   file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
   template_file_var.set(file_path)
def browse_target_file():
   file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
   target_file_var.set(file_path)
def browse_output_file():
   file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
   output_file_var.set(file_path)
def process_files():
   template_doc_path = template_file_var.get()
   target_doc_path = target_file_var.get()
   output_doc_path = output_file_var.get()
   if not template_doc_path or not target_doc_path or not output_doc_path:
       messagebox.showerror("Error", "请确保已选择所有文件。")
       return
   template_doc = docx.Document(template_doc_path)
   target_doc = docx.Document(target_doc_path)
   apply_template_formatting(template_doc, target_doc)
   target_doc.save(output_doc_path)
   messagebox.showinfo("Success", "文件已成功处理。")
root = tk.Tk()
root.title("Word格式应用器 by shuukou")
template_file_var = tk.StringVar()
target_file_var = tk.StringVar()
output_file_var = tk.StringVar()
tk.Label(root, text="模板文件：").grid(row=0, column=0, sticky="e")
tk.Entry(root, textvariable=template_file_var, width=40).grid(row=0, column=1)
tk.Button(root, text="浏览", command=browse_template_file).grid(row=0, column=2)
tk.Label(root, text="目标文件：").grid(row=1, column=0, sticky="e")
tk.Entry(root, textvariable=target_file_var,

width=40).grid(row=1, column=1)
tk.Button(root, text="浏览", command=browse_target_file).grid(row=1, column=2)
tk.Label(root, text="输出文件：").grid(row=2, column=0, sticky="e")
tk.Entry(root, textvariable=output_file_var, width=40).grid(row=2, column=1)
tk.Button(root, text="浏览", command=browse_output_file).grid(row=2, column=2)
tk.Button(root, text="开始处理", command=process_files).grid(row=3, column=1, pady=10)
root.mainloop()
